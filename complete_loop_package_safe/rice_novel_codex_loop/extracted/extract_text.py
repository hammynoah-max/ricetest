from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "source"
WORKBOOK_DIR = ROOT / "workbook"
OUT_DIR = ROOT / "extracted"


def clean_text(value: str) -> str:
    value = value.replace("\r", "\n")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def write_markdown(path: Path, title: str, body: str) -> None:
    path.write_text(f"# {title}\n\n{clean_text(body)}\n", encoding="utf-8")


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    chunks: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append(text)
    for table_index, table in enumerate(doc.tables, start=1):
        chunks.append(f"\n[Table {table_index}]")
        for row in table.rows:
            cells = [clean_text(cell.text).replace("\n", " / ") for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return "\n\n".join(chunks)


def extract_pdf(path: Path) -> str:
    import pdfplumber

    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            chunks.append(f"## Page {i}\n\n{text.strip()}")
    return "\n\n".join(chunks)


def xml_text(xml_bytes: bytes) -> str:
    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return ""
    parts: list[str] = []
    for elem in root.iter():
        tag = elem.tag.rsplit("}", 1)[-1]
        if tag in {"t", "text"} and elem.text:
            parts.append(elem.text)
        elif elem.text and elem.text.strip() and tag not in {"script", "style"}:
            text = elem.text.strip()
            if len(text) > 1:
                parts.append(text)
    return clean_text("\n".join(parts))


def extract_hwpx(path: Path) -> str:
    chunks: list[str] = []
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        section_names = [
            n for n in names
            if n.lower().endswith(".xml")
            and ("contents/" in n.lower() or "section" in n.lower())
        ]
        if not section_names:
            section_names = [n for n in names if n.lower().endswith(".xml")]
        for name in sorted(section_names):
            text = xml_text(zf.read(name))
            if text:
                chunks.append(f"## {name}\n\n{text}")
    return "\n\n".join(chunks)


def detect_kind(path: Path) -> str:
    suffix = path.suffix.lower()
    first = path.read_bytes()[:8]
    if suffix == ".pdf" or first.startswith(b"%PDF"):
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix == ".hwpx":
        return "hwpx"
    if first.startswith(b"PK"):
        try:
            with zipfile.ZipFile(path) as zf:
                names = "\n".join(zf.namelist()).lower()
            if "contents/" in names or "mimetype" in names:
                return "hwpx"
            if "word/document.xml" in names:
                return "docx"
        except zipfile.BadZipFile:
            pass
    return "unknown"


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    inventory = []
    failures = []

    files = sorted(SOURCE_DIR.iterdir()) + sorted(WORKBOOK_DIR.iterdir())
    for path in files:
        if not path.is_file():
            continue
        kind = detect_kind(path)
        out_name = f"{path.stem}_{kind}.md"
        out_path = OUT_DIR / out_name
        try:
            if kind == "pdf":
                body = extract_pdf(path)
            elif kind == "docx":
                body = extract_docx(path)
            elif kind == "hwpx":
                body = extract_hwpx(path)
            else:
                raise RuntimeError("unsupported file type")
            if not body.strip():
                raise RuntimeError("no text extracted")
            write_markdown(out_path, f"{path.name} ({kind})", body)
            inventory.append({
                "source": str(path.relative_to(ROOT)),
                "kind": kind,
                "output": str(out_path.relative_to(ROOT)),
                "status": "ok",
            })
        except Exception as exc:
            failures.append({
                "source": str(path.relative_to(ROOT)),
                "kind": kind,
                "error": str(exc),
            })

    (OUT_DIR / "extraction_inventory.json").write_text(
        json.dumps({"files": inventory, "failures": failures}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    lines = ["# 추출 결과 목록", ""]
    lines.append("## 성공")
    for item in inventory:
        lines.append(f"- `{item['source']}` -> `{item['output']}` ({item['kind']})")
    lines.append("")
    lines.append("## 실패")
    if failures:
        for item in failures:
            lines.append(f"- `{item['source']}` ({item['kind']}): {item['error']}")
    else:
        lines.append("- 없음")
    (OUT_DIR / "extraction_report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
